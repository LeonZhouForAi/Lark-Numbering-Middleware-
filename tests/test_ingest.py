import tempfile
import unittest
from pathlib import Path

from docx import Document

from feishu_rag.ingest import UnsupportedFileError, extract_sections, index_directory
from feishu_rag.store import IndexStore


class IngestTests(unittest.TestCase):
    def test_index_file_can_use_semantic_planner_metadata(self):
        class Planner:
            def plan(self, units):
                return {
                    "groups": [
                        {
                            "unit_ids": [unit.unit_id for unit in units],
                            "title": "付款语义标题",
                            "keywords": ["付款关键词"],
                            "summary": "付款语义摘要",
                        }
                    ]
                }

        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            docs = root / "docs"
            docs.mkdir()
            path = docs / "policy.txt"
            path.write_text("付款申请应由主管审批。", encoding="utf-8")
            store = IndexStore(root / "rag.sqlite3")
            try:
                count = index_directory(docs, store, semantic_planner=Planner())
                self.assertEqual(count, 1)
                result = store.search("付款关键词")[0].chunk
                self.assertEqual(result.content, "付款申请应由主管审批。")
            finally:
                store.close()

    def test_extracts_text_and_docx_sections(self):
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            txt = root / "制度.txt"
            txt.write_text("报销制度\n\n需要发票。", encoding="utf-8")
            docx_path = root / "流程.docx"
            doc = Document()
            doc.add_paragraph("付款申请流程")
            doc.add_paragraph("申请人提交审批。")
            doc.save(docx_path)

            txt_sections = extract_sections(txt)
            docx_sections = extract_sections(docx_path)

            self.assertIn("报销制度", "\n".join(section.text for section in txt_sections))
            self.assertIn("付款申请流程", "\n".join(section.text for section in docx_sections))

    def test_rejects_legacy_word_and_wps(self):
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            for suffix in (".doc", ".wps"):
                path = root / f"legacy{suffix}"
                path.write_bytes(b"legacy")
                with self.assertRaisesRegex(UnsupportedFileError, "先转换为 DOCX"):
                    extract_sections(path)

    def test_indexes_supported_files_recursively(self):
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            (root / "财务").mkdir()
            (root / "财务" / "报销.md").write_text("报销申请需要发票。", encoding="utf-8")
            store = IndexStore(root / "index.sqlite3")
            try:
                indexed = index_directory(root, store, max_chars=100)
                self.assertEqual(indexed, 1)
                self.assertEqual(store.count_documents(), 1)
                self.assertEqual(store.search("发票报销")[0].chunk.title, "报销")
            finally:
                store.close()


if __name__ == "__main__":
    unittest.main()
